from rest_framework import viewsets
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView
from .filters import UserInfoFilter
from django_filters.rest_framework import DjangoFilterBackend
from .pagination import UserInfoPagination
from collections import defaultdict
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from django.db import connection
from .utils import (
    degree_choices,
    job_classification_choices,
    classification_choices,
    vacation_choices,
    months_map,
    reversed_month_map,
)
from .models import (
    UserInfo,
    ShootingRecord,
    FitnessRecord,
    EducationRecord,
    LeaderOpinionRecord,
    GiftRecord,
    DisciplineRecord,
    VacationNote,
    EltemasRecord,
    Course,
    PreviousPosition,
    PreviousUnit,
    MedicalSituation,
)
from .serializers import (
    UserInfoSerializer,
    ShootingRecordSerializer,
    FitnessRecordSerializer,
    EducationRecordSerializer,
    LeaderOpinionRecordSerializer,
    GiftRecordSerializer,
    DisciplineRecordSerializer,
    VacationNoteSerializer,
    EltemasRecordSerializer,
    CourseSerializer,
    PreviousPositionSerializer,
    PreviousUnitSerializer,
    UsersNameSerializer,
    MedicalSituationSerializer,
    UserInfoListSerializer,
)


class UserInfoViewSet(viewsets.ModelViewSet):
    queryset = UserInfo.objects.prefetch_related(
        "shooting_records",
        "fitness_records",
        "education_records",
        "leader_opinion_records",
        "gift_records",
        "discipline_records",
        "eltemas_records",
        "courses",
        "previous_positions",
        "previous_units",
        "user_results",
        "medical_situations",
    )
    filter_backends = (DjangoFilterBackend,)
    filterset_class = UserInfoFilter
    pagination_class = UserInfoPagination
    serializer_class = UserInfoSerializer
    permission_classes = [IsAuthenticated]

    def get_serializer_class(self):
        if self.action == "list":
            return UserInfoListSerializer
        return UserInfoSerializer

    def get_queryset(self):
        qs = super().get_queryset()
        if self.action == "list":
            return qs.only(
                "id",
                "name",
                "job_classification",
                "classification",
                "militeryNumber",
                "vacation",
                "getInDate",
                "getOutDate",
                "dgree",
                "rank",
                "fasela",
                "saria",
                "is_soldier",
                "is_officer",
                "seniority",
                "batch",
                "unified_number",
            )
        return qs


class ShootingRecordViewSet(viewsets.ModelViewSet):
    queryset = ShootingRecord.objects.all()
    serializer_class = ShootingRecordSerializer
    permission_classes = [IsAuthenticated]


class FitnessRecordViewSet(viewsets.ModelViewSet):
    queryset = FitnessRecord.objects.all()
    serializer_class = FitnessRecordSerializer
    permission_classes = [IsAuthenticated]


class EducationRecordViewSet(viewsets.ModelViewSet):
    queryset = EducationRecord.objects.all()
    serializer_class = EducationRecordSerializer
    permission_classes = [IsAuthenticated]


class LeaderOpinionRecordViewSet(viewsets.ModelViewSet):
    queryset = LeaderOpinionRecord.objects.all()
    serializer_class = LeaderOpinionRecordSerializer
    permission_classes = [IsAuthenticated]


class GiftRecordViewSet(viewsets.ModelViewSet):
    queryset = GiftRecord.objects.all()
    serializer_class = GiftRecordSerializer
    permission_classes = [IsAuthenticated]


class DisciplineRecordViewSet(viewsets.ModelViewSet):
    queryset = DisciplineRecord.objects.all()
    serializer_class = DisciplineRecordSerializer
    permission_classes = [IsAuthenticated]


class VacationNotesViewSet(viewsets.ModelViewSet):
    queryset = VacationNote.objects.all()
    serializer_class = VacationNoteSerializer
    permission_classes = [IsAuthenticated]


class EltemasRecordViewSet(viewsets.ModelViewSet):
    queryset = EltemasRecord.objects.all()
    serializer_class = EltemasRecordSerializer
    permission_classes = [IsAuthenticated]


class CourseViewSet(viewsets.ModelViewSet):
    queryset = Course.objects.all()
    serializer_class = CourseSerializer
    # permission_classes = [IsAuthenticated]


class PreviousPositionViewSet(viewsets.ModelViewSet):
    queryset = PreviousPosition.objects.all()
    serializer_class = PreviousPositionSerializer
    permission_classes = [IsAuthenticated]


class PreviousUnitViewSet(viewsets.ModelViewSet):
    queryset = PreviousUnit.objects.all()
    serializer_class = PreviousUnitSerializer
    permission_classes = [IsAuthenticated]


class MedicalSituationViewSet(viewsets.ModelViewSet):
    queryset = MedicalSituation.objects.all()
    serializer_class = MedicalSituationSerializer
    permission_classes = [IsAuthenticated]


class UsersName(APIView):
    def get(self, request, *args, **kwargs):
        users = UserInfo.objects.filter(is_soldier=0)
        return Response(UsersNameSerializer(users, many=True).data)


class Statistics(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        from_month = request.GET.get("from")
        to_month = request.GET.get("to")
        cur_year = int(request.GET.get("year"))

        users = UserInfo.objects.all()
        sports = FitnessRecord.objects.all()
        educations = EducationRecord.objects.all()

        # Initialize counters
        officers = 0
        job_classification_map = defaultdict(int)
        classification_map = defaultdict(int)
        degree_map = defaultdict(int)
        rank_map = defaultdict(int)
        vacation_map = defaultdict(int)
        getIn_map = defaultdict(int)
        getOut_map = defaultdict(int)

        # User-related data processing
        for user in users:
            if user.is_officer:
                officers += 1

            degree_map[user.dgree] += 1
            rank_map[user.rank] += 1

            if not user.is_soldier:
                continue

            vacation_map[user.vacation] += 1
            job_classification_map[user.job_classification] += 1
            classification_map[user.classification] += 1
            getIn_map[str(user.getInDate)] += 1
            getOut_map[str(user.getOutDate)] += 1

        # Helper for processing monthly records
        def process_monthly_records(records):
            monthly_map = defaultdict(list)
            for record in records:
                if record.percentage == 0 or record.date.year != cur_year:
                    continue
                monthly_map[reversed_month_map[record.date.month]].append(
                    record.percentage
                )

            # Average and sort monthly data
            averaged_map = {
                month: "{0:.2f}".format(sum(values) / len(values))
                for month, values in monthly_map.items()
            }
            return sorted(
                [
                    (month, value)
                    for month, value in averaged_map.items()
                    if months_map[from_month]
                    <= months_map[month]
                    <= months_map[to_month]
                ],
                key=lambda x: months_map[x[0]],
            )

        # Process sports and education records
        sorted_sports_map = process_monthly_records(sports)
        sorted_educations_map = process_monthly_records(educations)

        # Fill missing keys with zero
        def fill_missing_keys(data_map, choices):
            for choice in choices:
                data_map.setdefault(choice, 0)

        fill_missing_keys(job_classification_map, job_classification_choices)
        fill_missing_keys(classification_map, classification_choices)
        fill_missing_keys(degree_map, degree_choices)
        fill_missing_keys(vacation_map, vacation_choices)

        # Clean unwanted entries
        degree_map.pop("لايكن", None)
        rank_map.pop("لايكن", None)

        # Prepare response data
        data = {
            "total": len(users),
            "soliders": degree_map.get("جندي", 0),
            "job_classification": dict(job_classification_map),
            "classification": dict(classification_map),
            "dgree": dict(degree_map),
            "rank": dict(rank_map),
            "vacation": dict(vacation_map),
            "getInDate": dict(sorted(getIn_map.items())),
            "getOutDate": dict(sorted(getOut_map.items())),
            "sports": {
                "keys": [x for x, y in sorted_sports_map],
                "values": [y for x, y in sorted_sports_map],
            },
            "educations": {
                "keys": [x for x, y in sorted_educations_map],
                "values": [y for x, y in sorted_educations_map],
            },
            "officers": officers,
        }
        return Response(data)


class DeleteRecords(APIView):
    queryset = UserInfo.objects.all()
    serializer_class = UserInfoSerializer
    permission_classes = [IsAuthenticated]

    def post(self, request, *args, **kwargs):
        users = request.data["users"]
        for user in users:
            if user != -1:
                obj = self.queryset.get(pk=user)
                obj.delete()
        return Response("ok")


class Vacation(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        vacation = self.request.GET.get("vacation")
        data = {
            "vacation_records": [],
        }

        # Default record structures for missing entries
        default_fitness_record = {
            "id": None,
            "running_rate": 0,
            "push_up": 0,
            "pull_up": 0,
            "crunch": 0,
            "running": 0,
            "percentage": "0.00",
            "weight": 0.0,
            "height": 0.0,
            "notes": "غير متوفر",
            "date": None,
            "date_modified": None,
            "date_created": None,
            "user": None,
        }

        default_education_record = {
            "id": None,
            "electronic_war": 0,
            "weapons": 0,
            "military_security": 0,
            "topography": 0,
            "chemical_war": 0,
            "presentation_of_self_mission": 0,
            "percentage": "0.00",
            "notes": "غير متوفر",
            "date": None,
            "date_created": None,
            "date_modified": None,
            "user": None,
        }

        default_shooting_record = {
            "id": None,
            "result": "غير متوفر",
            "data_modified": None,
            "date_created": None,
            "notes": "غير متوفر",
            "date": None,
            "user": None,
        }

        # Fetch users based on vacation filter and prefetch related records
        users = UserInfo.objects.prefetch_related(
            "shooting_records", "fitness_records", "education_records"
        ).filter(vacation=vacation)

        # Process each user
        for user in users:
            if not user.is_soldier:
                continue
            # Get related records
            fitness_records = FitnessRecordSerializer(
                user.fitness_records.all()[:2], many=True
            ).data
            education_records = EducationRecordSerializer(
                user.education_records.all()[:2], many=True
            ).data
            shooting_records = ShootingRecordSerializer(
                user.shooting_records.all()[:2], many=True
            ).data

            # Add default records if necessary
            while len(fitness_records) < 2:
                fitness_records.append(default_fitness_record)

            while len(education_records) < 2:
                education_records.append(default_education_record)

            while len(shooting_records) < 2:
                shooting_records.append(default_shooting_record)

            data["vacation_records"].append(
                {
                    "id": user.id,
                    "name": user.name,
                    "fullAddress": user.fullAddress,
                    "job_classification": user.job_classification,
                    "vacation": user.vacation,
                    "dgree": user.dgree,
                    "first_record": {
                        "fitness_percentage": fitness_records[0]["percentage"],
                        "education_percentage": education_records[0]["percentage"],
                        "shooting_value": shooting_records[0]["result"],
                        "running_value": fitness_records[0]["running"],
                        "notes": fitness_records[0]["notes"],
                        "weight": fitness_records[0]["weight"],
                        "height": fitness_records[0]["height"],
                    },
                    "second_record": {
                        "fitness_percentage": fitness_records[1]["percentage"],
                        "education_percentage": education_records[1]["percentage"],
                        "shooting_value": shooting_records[1]["result"],
                        "running_value": fitness_records[1]["running"],
                        "notes": fitness_records[1]["notes"],
                        "weight": fitness_records[1]["weight"],
                        "height": fitness_records[1]["height"],
                    },
                    "fitness_records": fitness_records,
                    "education_records": education_records,
                    "shooting_records": shooting_records,
                }
            )

        data["vacation_records"].sort(
            key=lambda item: float(item["first_record"]["fitness_percentage"]),
            reverse=True,
        )

        return Response(data)


class FitnessShow(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        data = {
            "records": [],
        }

        default_fitness_record = {
            "id": None,
            "running_rate": 0,
            "push_up": 0,
            "pull_up": 0,
            "crunch": 0,
            "running": 0,
            "percentage": "0.00",
            "weight": 0.0,
            "height": 0.0,
            "notes": "غير متوفر",
            "date": None,
            "date_modified": None,
            "date_created": None,
            "user": None,
        }

        users = UserInfo.objects.prefetch_related("fitness_records")

        for user in users:
            if not user.is_soldier:
                continue
            fitness_records = FitnessRecordSerializer(
                user.fitness_records.all()[:1], many=True
            ).data

            while len(fitness_records) < 1:
                fitness_records.append(default_fitness_record)

            data["records"].append(
                {
                    "id": user.id,
                    "name": user.name,
                    "vacation": user.vacation,
                    "is_officer": user.is_officer,
                    "is_soldier": user.is_soldier,
                    "fitness_records": fitness_records,
                }
            )
        return Response(data)


class EducationsShow(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        data = {
            "records": [],
        }

        default_education_record = {
            "id": None,
            "electronic_war": 0,
            "weapons": 0,
            "military_security": 0,
            "topography": 0,
            "chemical_war": 0,
            "presentation_of_self_mission": 0,
            "percentage": "0.00",
            "notes": "غير متوفر",
            "date": None,
            "date_created": None,
            "date_modified": None,
            "user": None,
        }

        users = UserInfo.objects.prefetch_related("education_records")

        for user in users:
            if not user.is_soldier:
                continue
            education_records = EducationRecordSerializer(
                user.education_records.all()[:1], many=True
            ).data

            while len(education_records) < 1:
                education_records.append(default_education_record)

            data["records"].append(
                {
                    "id": user.id,
                    "name": user.name,
                    "vacation": user.vacation,
                    "is_officer": user.is_officer,
                    "is_soldier": user.is_soldier,
                    "education_records": education_records,
                }
            )
        return Response(data)


class Reports(APIView):
    permission_classes = [
        IsAuthenticated,
    ]

    def post(self, request, *args, **kwargs):
        # Fetch users with is_soldier=1
        vacation = request.data["vacation"]
        users = UserInfo.objects.filter(is_soldier=1, vacation=vacation)
        try:
            # Load the Word documents
            sport_doc = Document("media/sport.docx")
            education_doc = Document("media/education.docx")
            shooting_doc = Document("media/shooting.docx")
            running_doc = Document("media/running.docx")
        except FileNotFoundError as e:
            return Response("Files not found", status=404)
        # Access the tables in each document
        sport_table = sport_doc.tables[0]
        education_table = education_doc.tables[0]
        shooting_table = shooting_doc.tables[0]
        running_table = running_doc.tables[0]

        # Iterate over users and update all tables
        for user in users:
            # Update sport.docx
            new_row = sport_table.add_row()
            cells = new_row.cells
            cells[4].text = user.name
            run = cells[4].paragraphs[0].runs[0]
            run.font.size = Pt(14)
            cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Update education.docx
            new_row = education_table.add_row()
            cells = new_row.cells
            cells[5].text = user.name
            run = cells[5].paragraphs[0].runs[0]
            run.font.size = Pt(14)
            cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Update shooting.docx
            new_row = shooting_table.add_row()
            cells = new_row.cells
            cells[1].text = user.name
            run = cells[1].paragraphs[0].runs[0]
            run.font.size = Pt(14)
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Update running.docx
            new_row = running_table.add_row()
            cells = new_row.cells
            cells[3].text = user.name
            run = cells[3].paragraphs[0].runs[0]
            run.font.size = Pt(14)
            cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Save the updated documents
        sport_doc_path = "media/كشف اللياقة.docx"
        education_doc_path = "media/كشف التعليم.docx"
        shooting_doc_path = "media/كشف الرماية.docx"
        running_doc_path = "media/كشف الضاحية.docx"

        sport_doc.save(sport_doc_path)
        education_doc.save(education_doc_path)
        shooting_doc.save(shooting_doc_path)
        running_doc.save(running_doc_path)

        base_url = request.build_absolute_uri("/")
        sport_doc_url = f"{base_url}media/كشف اللياقة.docx"
        education_doc_url = f"{base_url}media/كشف التعليم.docx"
        shooting_doc_url = f"{base_url}media/كشف الرماية.docx"
        running_doc_url = f"{base_url}media/كشف الضاحية.docx"

        return Response(
            {
                "message": "Documents updated successfully.",
                "links": {
                    "sport": sport_doc_url,
                    "education": education_doc_url,
                    "shooting": shooting_doc_url,
                    "running": running_doc_url,
                },
            }
        )


class Tesing(APIView):
    # permission_classes = [
    #     IsAuthenticated,
    # ]

    def post(self, request, *args, **kwargs):
        print(connection.queries)
        return Response("ok")
